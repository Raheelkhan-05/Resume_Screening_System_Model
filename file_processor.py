"""
File processing module for extracting text from PDF and DOCX files
"""

import logging
from typing import Optional, Dict, Any
import PyPDF2
from docx import Document
import re
from pathlib import Path
from utils import sanitize_text

logger = logging.getLogger(__name__)

class FileProcessor:
    """Handles extraction of text from various file formats"""
    
    def __init__(self):
        self.supported_types = {
            '.pdf': self._extract_from_pdf,
            '.docx': self._extract_from_docx,
            '.txt': self._extract_from_txt
        }
    
    def extract_text_from_file(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from file and return structured result
        
        Args:
            file_path: Path to the file
            
        Returns:
            Dict containing extracted text and metadata
        """
        try:
            if not Path(file_path).exists():
                raise FileNotFoundError(f"File not found: {file_path}")
            
            file_ext = Path(file_path).suffix.lower()
            
            if file_ext not in self.supported_types:
                raise ValueError(f"Unsupported file type: {file_ext}")
            
            # Extract text using appropriate method
            raw_text = self.supported_types[file_ext](file_path)
            
            # Clean and sanitize text
            cleaned_text = sanitize_text(raw_text)
            
            # Basic statistics
            word_count = len(cleaned_text.split())
            char_count = len(cleaned_text)
            
            result = {
                'file_path': file_path,
                'file_type': file_ext,
                'raw_text': raw_text,
                'cleaned_text': cleaned_text,
                'word_count': word_count,
                'char_count': char_count,
                'extraction_successful': True,
                'error_message': None
            }
            
            logger.info(f"Successfully extracted {word_count} words from {file_path}")
            return result
            
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return {
                'file_path': file_path,
                'file_type': file_ext if 'file_ext' in locals() else 'unknown',
                'raw_text': '',
                'cleaned_text': '',
                'word_count': 0,
                'char_count': 0,
                'extraction_successful': False,
                'error_message': str(e)
            }
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"
                    
        except Exception as e:
            logger.error(f"Error reading PDF {file_path}: {str(e)}")
            raise
            
        return text
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
                        
        except Exception as e:
            logger.error(f"Error reading DOCX {file_path}: {str(e)}")
            raise
            
        return text
    
    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                text = file.read()
        except Exception as e:
            logger.error(f"Error reading TXT {file_path}: {str(e)}")
            raise
            
        return text

class TextPreprocessor:
    """Handles text preprocessing and normalization"""
    
    def __init__(self, config: Optional[Dict] = None):
        self.config = config or {
            'remove_stopwords': True,
            'lemmatize': True,
            'lowercase': True,
            'normalize_whitespace': True
        }
        self._setup_nlp()
    
    def _setup_nlp(self):
        """Setup NLP tools"""
        try:
            import nltk
            import spacy
            
            # Download NLTK data if not present
            try:
                nltk.data.find('tokenizers/punkt')
            except LookupError:
                nltk.download('punkt')
                
            try:
                nltk.data.find('corpora/stopwords')
            except LookupError:
                nltk.download('stopwords')
                
            try:
                nltk.data.find('corpora/wordnet')
            except LookupError:
                nltk.download('wordnet')
            
            from nltk.corpus import stopwords
            from nltk.stem import WordNetLemmatizer
            
            self.stop_words = set(stopwords.words('english'))
            self.lemmatizer = WordNetLemmatizer()
            
            logger.info("NLP tools initialized successfully")
            
        except Exception as e:
            logger.warning(f"Could not initialize advanced NLP tools: {e}")
            self.stop_words = set()
            self.lemmatizer = None
    
    def preprocess_text(self, text: str) -> Dict[str, Any]:
        """
        Preprocess and clean text
        
        Args:
            text: Raw text to preprocess
            
        Returns:
            Dict containing processed text and metadata
        """
        if not text:
            return {
                'original_text': text,
                'processed_text': '',
                'preprocessing_applied': []
            }
        
        processed_text = text
        applied_steps = []
        
        # Lowercase
        if self.config.get('lowercase', True):
            processed_text = processed_text.lower()
            applied_steps.append('lowercase')
        
        # Normalize whitespace
        if self.config.get('normalize_whitespace', True):
            processed_text = re.sub(r'\s+', ' ', processed_text).strip()
            applied_steps.append('normalize_whitespace')
        
        # Remove special characters (keep alphanumeric and basic punctuation)
        processed_text = re.sub(r'[^\w\s\.,;:\-\(\)]+', ' ', processed_text)
        applied_steps.append('remove_special_chars')
        
        # Remove stopwords and lemmatize if tools are available
        if self.config.get('remove_stopwords', True) and self.stop_words:
            words = processed_text.split()
            
            # Remove stopwords
            words = [word for word in words if word not in self.stop_words]
            applied_steps.append('remove_stopwords')
            
            # Lemmatize
            if self.config.get('lemmatize', True) and self.lemmatizer:
                words = [self.lemmatizer.lemmatize(word) for word in words]
                applied_steps.append('lemmatize')
            
            processed_text = ' '.join(words)
        
        return {
            'original_text': text,
            'processed_text': processed_text,
            'preprocessing_applied': applied_steps,
            'word_count_before': len(text.split()),
            'word_count_after': len(processed_text.split())
        }